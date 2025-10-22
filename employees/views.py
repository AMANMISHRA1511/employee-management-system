import os
import shutil
import random
import string
import json
import traceback
from datetime import datetime
from io import BytesIO
import re
from django.contrib.auth.models import User
from django.contrib.auth import get_user_model
from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse, HttpResponse
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.views.decorators.csrf import csrf_exempt
from django.core.paginator import Paginator
from django.db.models import Q
from django.utils import timezone
from django.core.management import call_command
from django.conf import settings
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from .models import Employee
from .forms import EmployeeForm

User = get_user_model()


# ------------------------------
# Utility: Generate unique username
# ------------------------------
def generate_unique_username(base_username):
    base_username = base_username.lower().replace(" ", "")
    for _ in range(1000):
        random_number = random.randint(100, 999)
        username = f"{base_username}{random_number}"
        if not User.objects.filter(username=username).exists():
            return username
    counter = 1
    while True:
        username = f"{base_username}{counter}"
        if not User.objects.filter(username=username).exists():
            return username
        counter += 1


# ------------------------------
# Dashboard
# ------------------------------
@login_required
def dashboard(request):
    employees = Employee.objects.all() if request.user.user_type == 'admin' else Employee.objects.filter(user=request.user)
    total_employees = employees.count()
    active_employees = employees.filter(user__is_active=True).count()
    departments = employees.values('department').distinct().count()
    current_month = timezone.now().month
    current_year = timezone.now().year
    new_employees = employees.filter(join_date__month=current_month, join_date__year=current_year).count()
    recent_employees = employees.order_by('-join_date')[:5]
    dept_list = employees.values_list('department', flat=True).distinct()
    dept_stats = {dept: employees.filter(department=dept).count() for dept in dept_list}

    context = {
        'total_employees': total_employees,
        'active_employees': active_employees,
        'departments': departments,
        'new_employees': new_employees,
        'recent_employees': recent_employees,
        'dept_stats': dept_stats,
    }
    return render(request, 'employees/dashboard.html', context)


# ------------------------------
# Employee List
# ------------------------------
@login_required
def employee_list(request):
    if request.user.user_type != 'admin':
        messages.error(request, "You don't have permission to view this page.")
        return redirect('dashboard')

    employees = Employee.objects.all()

    # Filtering
    department_filter = request.GET.get('department', '')
    status_filter = request.GET.get('status', '')
    if department_filter:
        employees = employees.filter(department__iexact=department_filter)
    if status_filter:
        employees = employees.filter(status__iexact=status_filter)

    # Searching
    search_query = request.GET.get('search', '')
    if search_query:
        employees = employees.filter(
            Q(user__first_name__icontains=search_query) |
            Q(user__last_name__icontains=search_query) |
            Q(user__email__icontains=search_query) |
            Q(department__icontains=search_query) |
            Q(position__icontains=search_query)
        )

    # Sorting
    sort_by = request.GET.get('sort', 'user__first_name')
    valid_sort_fields = ['user__first_name', 'department', 'join_date', 'status']
    if sort_by not in valid_sort_fields:
        sort_by = 'user__first_name'
    employees = employees.order_by(sort_by)

    # Pagination
    paginator = Paginator(employees, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Flags for template filters
    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'is_hr': department_filter == 'HR',
        'is_it': department_filter == 'IT',
        'is_finance': department_filter == 'Finance',
        'is_marketing': department_filter == 'Marketing',
        'is_sales': department_filter == 'Sales',
        'is_operations': department_filter == 'Operations',
        'is_active': status_filter == 'active',
        'is_inactive': status_filter == 'inactive',
        'is_sort_name': sort_by == 'user__first_name',
        'is_sort_department': sort_by == 'department',
        'is_sort_join_date': sort_by == 'join_date',
        'is_sort_status': sort_by == 'status',
    }

    return render(request, 'employees/employee_list.html', context)

# ------------------------------
# Add Employee
# ------------------------------
@login_required
def add_employee(request):
    if request.user.user_type != 'admin':
        messages.error(request, "You don't have permission to add employees.")
        return redirect('dashboard')

    if request.method == 'POST':
        form = EmployeeForm(request.POST, request.FILES)
        if form.is_valid():
            base_username = re.sub(r'[^a-zA-Z0-9]', '', form.cleaned_data['email'].split('@')[0])
            username = generate_unique_username(base_username)
            temp_password = ''.join(random.choices(string.ascii_letters + string.digits, k=8))

            user = User.objects.create_user(
                username=username,
                email=form.cleaned_data['email'],
                password=temp_password,
                first_name=form.cleaned_data['first_name'],
                last_name=form.cleaned_data['last_name'],
                user_type='employee'
            )

            employee = form.save(commit=False)
            employee.user = user
            employee.save()

            messages.success(request, f'Employee added successfully! Username: {username} | Temporary Password: {temp_password}')
            return redirect('employee_list')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = EmployeeForm()

    return render(request, 'employees/add_employee.html', {'form': form})


# ------------------------------
# Edit Employee
# ------------------------------
@login_required
def edit_employee(request, pk):
    if request.user.user_type != 'admin':
        messages.error(request, "You don't have permission to edit employees.")
        return redirect('dashboard')

    employee = get_object_or_404(Employee, pk=pk)

    # Handle missing linked user safely
    try:
        user = employee.user
    except User.DoesNotExist:
        messages.error(request, "Cannot edit this employee because the linked user is missing.")
        return redirect('employee_list')

    if request.method == 'POST':
        form = EmployeeForm(request.POST, request.FILES, instance=employee)
        if form.is_valid():
            form.save()
            messages.success(request, 'Employee updated successfully!')
            return redirect('employee_list')
    else:
        form = EmployeeForm(instance=employee)

    return render(request, 'employees/add_employee.html', {
        'form': form,
        'editing': True,
        'employee': employee
    })


# ------------------------------
# Delete Employee
@login_required
def delete_employee(request, employee_id):
    # Only admin users can delete
    if request.user.user_type != 'admin':
        messages.error(request, "You don't have permission to delete this employee.")
        return redirect('employee_list')

    # Get employee or 404
    employee = get_object_or_404(Employee, id=employee_id)

    if request.method == 'POST':
        # Delete employee first
        employee.delete()

        # Delete linked User if exists
        if hasattr(employee, 'user') and employee.user is not None:
            try:
                employee.user.delete()
            except User.DoesNotExist:
                pass  # User already deleted, ignore

        messages.success(request, f"Employee '{employee.full_name}' deleted successfully.")
        return redirect('employee_list')

    return render(request, 'employees/delete_employee.html', {'employee': employee})

# ------------------------------
@login_required
def export_employees(request, format):
    employees = Employee.objects.select_related('user').all()

    # CSV
    if format.lower() == 'csv':
        import csv
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="employees.csv"'
        writer = csv.writer(response)
        writer.writerow(['First Name', 'Last Name', 'Department', 'Email', 'Phone', 'Address', 'Join Date', 'Position', 'Status'])
        for e in employees:
            writer.writerow([
                e.user.first_name if e.user else 'N/A',
                e.user.last_name if e.user else 'N/A',
                e.department, e.email, e.phone or '', e.address or '',
                str(e.join_date), e.position, e.status
            ])
        return response

    # JSON
    elif format.lower() == 'json':
        data = []
        for e in employees:
            data.append({
                'first_name': e.user.first_name if e.user else 'N/A',
                'last_name': e.user.last_name if e.user else 'N/A',
                'department': e.department,
                'email': e.email,
                'phone': e.phone or '',
                'address': e.address or '',
                'join_date': str(e.join_date),
                'position': e.position,
                'status': e.status
            })
        return JsonResponse(data, safe=False)

    # Word
    elif format.lower() == 'word':
        document = Document()
        document.add_heading("Employee Report", level=1)
        table = document.add_table(rows=1, cols=9)
        table.style = 'LightShading-Accent1'
        headers = ['First Name', 'Last Name', 'Department', 'Email', 'Phone', 'Address', 'Join Date', 'Position', 'Status']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        for e in employees:
            row_cells = table.add_row().cells
            row_cells[0].text = e.user.first_name if e.user else 'N/A'
            row_cells[1].text = e.user.last_name if e.user else 'N/A'
            row_cells[2].text = e.department
            row_cells[3].text = e.email
            row_cells[4].text = e.phone or ''
            row_cells[5].text = e.address or ''
            row_cells[6].text = str(e.join_date)
            row_cells[7].text = e.position
            row_cells[8].text = e.status
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="employees.docx"'
        document.save(response)
        return response

    # PDF
    elif format.lower() == 'pdf':
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        styles = getSampleStyleSheet()
        elements.append(Paragraph("Employee Report", styles['Heading1']))
        data = [['First Name', 'Last Name', 'Department', 'Email', 'Phone', 'Address', 'Join Date', 'Position', 'Status']]
        for e in employees:
            data.append([
                e.user.first_name if e.user else 'N/A',
                e.user.last_name if e.user else 'N/A',
                e.department, e.email, e.phone or '', e.address or '',
                str(e.join_date), e.position, e.status
            ])
        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#d9d9d9')),
            ('TEXTCOLOR',(0,0),(-1,0),colors.black),
            ('ALIGN',(0,0),(-1,-1),'CENTER'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,0), 12),
            ('BOTTOMPADDING', (0,0), (-1,0), 8),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ]))
        elements.append(table)
        doc.build(elements)
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="employees.pdf"'
        response.write(buffer.getvalue())
        buffer.close()
        return response

    else:
        return HttpResponse("Unsupported format", status=400)


# ------------------------------
# Clear All Data
# ------------------------------
@login_required
def clear_all_data(request):
    if request.method != 'POST':
        messages.error(request, "Invalid request method.")
        return redirect('employee_list')

    try:
        employees = Employee.objects.select_related('user').all()
        for emp in employees:
            if emp.user:
                emp.user.delete()
        messages.success(request, "All employee data cleared successfully.")
        return redirect('employee_list')
    except Exception as e:
        traceback.print_exc()
        messages.error(request, f"Error clearing data: {str(e)}")
        return redirect('employee_list')


# ------------------------------
# Backup
# ------------------------------
@login_required
def create_backup(request):
    if request.user.user_type != 'admin':
        messages.error(request, "You don't have permission to create backups.")
        return redirect('dashboard')

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_dir = os.path.join(settings.BASE_DIR, "backups")
    os.makedirs(backup_dir, exist_ok=True)
    backup_filename = f"backup_{timestamp}.json"
    backup_path = os.path.join(backup_dir, backup_filename)

    with open(backup_path, "w", encoding="utf-8") as f:
        call_command("dumpdata", "--exclude", "sessions", "--indent", "2", stdout=f)

    with open(backup_path, "rb") as f:
        response = HttpResponse(f.read(), content_type="application/json")
        response["Content-Disposition"] = f'attachment; filename="{backup_filename}"'
        return response


@login_required
def create_full_backup(request):
    if request.user.user_type != 'admin':
        messages.error(request, "You don't have permission to create backups.")
        return redirect('dashboard')

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_dir = os.path.join(settings.BASE_DIR, "backups", f"backup_{timestamp}")
    os.makedirs(backup_dir, exist_ok=True)

    # Database backup
    db_file = os.path.join(backup_dir, "db_backup.json")
    with open(db_file, "w", encoding="utf-8") as f:
        call_command("dumpdata", "--exclude", "sessions", "--indent", "2", stdout=f)

    # Media backup
    media_dir = os.path.join(settings.BASE_DIR, "media")
    if os.path.exists(media_dir):
        shutil.copytree(media_dir, os.path.join(backup_dir, "media"))

    # Create ZIP
    zip_filename = f"backup_{timestamp}.zip"
    zip_path = os.path.join(settings.BASE_DIR, "backups", zip_filename)
    shutil.make_archive(zip_path.replace(".zip", ""), "zip", backup_dir)

    with open(zip_path, "rb") as f:
        response = HttpResponse(f.read(), content_type="application/zip")
        response["Content-Disposition"] = f'attachment; filename="{zip_filename}"'
        return response


# ------------------------------
# Simple Pages
# ------------------------------
@login_required
def files_page(request):
    if request.user.user_type != 'admin':
        messages.error(request, "You don't have permission to access this page.")
        return redirect('dashboard')
    return render(request, 'employees/files.html')


@login_required
def settings_page(request):
    return render(request, 'employees/settings.html')


@login_required
def employee_details(request, pk):
    employee = get_object_or_404(Employee, pk=pk)
    return render(request, 'employees/employee_details.html', {'employee': employee})


# ------------------------------
# Import Employees (JSON)
# ------------------------------
@login_required
@csrf_exempt
def import_employees(request):
    if request.user.user_type != 'admin':
        return JsonResponse({'error': 'Permission denied'}, status=403)

    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid request method'}, status=400)

    try:
        data = json.loads(request.body)
        count = 0
        for item in data:
            email = item.get('email')
            first_name = item.get('first_name')
            last_name = item.get('last_name')
            department = item.get('department', '')
            position = item.get('position', '')
            status = item.get('status', 'active')

            if not email:
                continue

            if Employee.objects.filter(user__email=email).exists():
                continue

            temp_password = 'Temp1234'
            username = generate_unique_username(email.split('@')[0])
            user = User.objects.create_user(
                username=username,
                email=email,
                password=temp_password,
                first_name=first_name,
                last_name=last_name,
                user_type='employee'
            )

            Employee.objects.create(
                user=user,
                department=department,
                position=position,
                status=status
            )
            count += 1

        return JsonResponse({'success': True, 'imported_count': count})
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
    
   
 
